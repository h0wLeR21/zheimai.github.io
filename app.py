import os
import glob
import json
from flask import Flask, render_template, request, jsonify, redirect, url_for, send_from_directory, session
from docx import Document
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import CharacterTextSplitter
from langchain_huggingface.embeddings import HuggingFaceEmbeddings
from langchain.chains import RetrievalQA
from langchain_groq import ChatGroq
from assessment_creator import AlzheimersAssessment
import traceback
import sqlite3
from datetime import datetime

def setup_reminder_database():
    """Set up the SQLite database for storing reminders"""
    db_path = "reminders.db"
    
    # Check if database already exists
    db_exists = os.path.exists(db_path)
    
    # Connect to database (will create it if it doesn't exist)
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    
    # Create reminders table if it doesn't exist
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS reminders (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        patient_name TEXT NOT NULL,
        title TEXT NOT NULL,
        description TEXT,
        time TEXT NOT NULL,
        frequency TEXT NOT NULL,
        created_by TEXT NOT NULL,
        active INTEGER DEFAULT 1,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    ''')
    
    conn.commit()
    conn.close()
    
    return db_exists

# Add this to your app.py to initialize the database when the app starts
def init_app():
    db_exists = setup_reminder_database()
    if not db_exists:
        print("Reminder database created successfully.")
    else:
        print("Reminder database already exists.")

# Call this function when your app starts


app = Flask(__name__, static_url_path='/static')
app.secret_key = os.urandom(24)  # For session management

# API Key (Use environment variable for security)
API_KEY = os.getenv("GROQ_API_KEY", "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")

# Define directory containing .docx files
docx_directory = "./data"
reports_directory = "./Reports"

# Ensure reports directory exists
os.makedirs(reports_directory, exist_ok=True)

# Initialize the QA chain
qa_chain = None

# System message defining the AI's behavior
SYSTEM_MESSAGE = (
    "Your name is Zheim. "
    "You are a friendly and intelligent AI assistant designed specifically for Alzheimer's patients, caretakers, and doctors. "
    "You can engage in casual discussions, answer general knowledge questions, and help with a wide range of topics. "
    "You must only use information from the provided documents when the user explicitly asks about them. "
    "When discussing Alzheimer's care, be empathetic, accurate, and supportive. "
    "Maintain a warm and engaging tone, but respect the user's privacy at all times. "
    "If the user asks an unclear or broad question, seek clarification before proceeding."
)

def initialize_qa_system():
    global qa_chain
    
    # Function to load and extract text from .docx files
    def load_docx_files(directory):
        if not os.path.exists(directory):
            os.makedirs(directory)
            return "No documents loaded yet."
        
        text_data = ""
        docx_files = glob.glob(os.path.join(directory, "*.docx"))

        if not docx_files:
            return "No .docx files found in the directory."

        for file_path in docx_files:
            try:
                doc = Document(file_path)
                for para in doc.paragraphs:
                    text_data += para.text + "\n"
            except Exception as e:
                print(f"Error reading {file_path}: {e}")
        
        return text_data

    try:
        # Extract text from docx files
        docx_text = load_docx_files(docx_directory)

        # Split text into smaller chunks for vector storage
        text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        texts = text_splitter.split_text(docx_text)

        # Create vector database using HuggingFace embeddings
        embeddings = HuggingFaceEmbeddings()
        db = FAISS.from_texts(texts, embeddings)
        retriever = db.as_retriever()

        # Initialize ChatGroq LLM
        llm = ChatGroq(
            api_key=API_KEY,
            model_name="llama3-70b-8192"
        )

        # Create retrieval-based QA chain
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm, 
            retriever=retriever
        )
        
        return True
    except Exception as e:
        print(f"Error initializing QA system: {e}")
        return False

# Function to get available questions from questions.docx
def get_assessment_questions():
    try:
        questions_path = "questions.docx"
        if not os.path.exists(questions_path):
            return []
        
        doc = Document(questions_path)
        questions = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        return questions
    except Exception as e:
        print(f"Error reading questions file: {e}")
        return []

# Function to get available assessment reports
def get_assessment_reports():
    try:
        if not os.path.exists(reports_directory):
            return []
        
        reports = glob.glob(os.path.join(reports_directory, "*.docx"))
        return [os.path.basename(report) for report in reports]
    except Exception as e:
        print(f"Error reading reports: {e}")
        return []

# Routes for serving static files
@app.route('/static/css/<path:filename>')
def serve_css(filename):
    return send_from_directory('static/css', filename)

@app.route('/static/js/<path:filename>')
def serve_js(filename):
    return send_from_directory('static/js', filename)

# Main routes
@app.route('/')
def index():
    return redirect(url_for('login'))

@app.route('/login')
@app.route('/login.html')
def login():
    return render_template('login.html')

@app.route('/home')
@app.route('/home.html')
def home():
    return render_template('home.html')

@app.route('/llm')
@app.route('/llm.html')
def llm():
    return render_template('llm.html')

# Assessment routes
@app.route('/assessment')
@app.route('/assessment.html')
def assessment():
    questions = get_assessment_questions()
    reports = get_assessment_reports()
    return render_template('assessment.html', questions=questions, reports=reports)

@app.route('/reports/<path:filename>')
def view_report(filename):
    return send_from_directory('Reports', filename)

@app.route('/create_assessment', methods=['POST'])
def create_assessment():
    try:
        data = request.json
        patient_name = data.get('patient_name', '')
        selected_questions = data.get('selected_questions', [])
        answers = data.get('answers', [])
        
        if not patient_name or not selected_questions or not answers:
            return jsonify({"error": "Missing required information"}), 400
        
        # Create assessment instance
        assessment = AlzheimersAssessment(patient_name)
        
        # Manually set the questions and answers
        assessment.selected_questions = selected_questions
        assessment.answers = answers
        
        # Generate report
        report_file = assessment.generate_report()
        
        return jsonify({
            "success": True,
            "message": "Assessment created successfully",
            "report_file": os.path.basename(report_file)
        })
    except Exception as e:
        print(f"Error creating assessment: {e}")
        return jsonify({"error": f"Error creating assessment: {str(e)}"}), 500
# Make sure you have these imports at the top of your app.py file

# Add this route to your app.py file
@app.route('/generate_memory_report', methods=['POST'])
def generate_memory_report():
    try:
        # Import here to avoid circular imports
        from report_assistant import get_all_assessment_reports, analyze_memory_trends
        
        # Ensure we have valid JSON
        if not request.is_json:
            return jsonify({"error": "Invalid request format. JSON required"}), 400
            
        data = request.json
        patient_name = data.get('patient_name', '')
        
        if not patient_name:
            return jsonify({"error": "Please enter a patient name"}), 400
        
        # Use your existing functions from report_assistant.py
        reports = get_all_assessment_reports(patient_name)
        
        if not reports:
            return jsonify({"error": f"No assessment reports found for {patient_name}"}), 404
        
        # Run the analysis using your existing function
        analysis = analyze_memory_trends(reports, patient_name)
        
        return jsonify({"success": True, "analysis": analysis})
    
    except Exception as e:
        # Print detailed error information to server console
        print(f"Error in generate_memory_report: {str(e)}")
        print(traceback.format_exc())
        return jsonify({"error": str(e)}), 500
# Add these routes to your app.py file
# Replace your existing get_reminders route with this updated version

@app.route('/get_reminders', methods=['POST'])
def get_reminders():
    try:
        data = request.json
        patient_name = data.get('patient_name', '')
        
        if not patient_name:
            return jsonify({"error": "Patient name is required"}), 400
        
        print(f"Fetching reminders for patient: {patient_name}")  # Debug log
        
        conn = sqlite3.connect('reminders.db')
        conn.row_factory = sqlite3.Row  # This enables column access by name
        cursor = conn.cursor()
        
        # Get active reminders for the patient
        cursor.execute('''
        SELECT id, patient_name, title, description, time, frequency, created_by, active, created_at
        FROM reminders
        WHERE patient_name = ? AND active = 1
        ORDER BY time ASC
        ''', (patient_name,))
        
        rows = cursor.fetchall()
        reminders = [{key: row[key] for key in row.keys()} for row in rows]
        
        print(f"Found {len(reminders)} reminders for {patient_name}")  # Debug log
        conn.close()
        
        return jsonify({"success": True, "reminders": reminders})
    
    except Exception as e:
        print(f"Error getting reminders: {str(e)}")
        import traceback
        traceback.print_exc()  # Print detailed error info
        return jsonify({"error": str(e)}), 500
    
@app.route('/add_reminder', methods=['POST'])
def add_reminder():
    try:
        data = request.json
        patient_name = data.get('patient_name', '')
        title = data.get('title', '')
        description = data.get('description', '')
        time = data.get('time', '')
        frequency = data.get('frequency', '')
        created_by = data.get('created_by', '')
        
        # Validate required fields
        if not all([patient_name, title, time, frequency, created_by]):
            return jsonify({"error": "Missing required fields"}), 400
        
        conn = sqlite3.connect('reminders.db')
        cursor = conn.cursor()
        
        # Insert new reminder
        cursor.execute('''
        INSERT INTO reminders (patient_name, title, description, time, frequency, created_by)
        VALUES (?, ?, ?, ?, ?, ?)
        ''', (patient_name, title, description, time, frequency, created_by))
        
        reminder_id = cursor.lastrowid
        conn.commit()
        conn.close()
        
        return jsonify({
            "success": True, 
            "message": "Reminder added successfully", 
            "reminder_id": reminder_id
        })
    
    except Exception as e:
        print(f"Error adding reminder: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/update_reminder', methods=['POST'])
def update_reminder():
    try:
        data = request.json
        reminder_id = data.get('id')
        
        if not reminder_id:
            return jsonify({"error": "Reminder ID is required"}), 400
        
        # Get fields to update
        fields = {}
        for field in ['title', 'description', 'time', 'frequency', 'active']:
            if field in data:
                fields[field] = data[field]
        
        if not fields:
            return jsonify({"error": "No fields to update"}), 400
        
        # Build the SQL query
        set_clause = ", ".join([f"{field} = ?" for field in fields])
        values = list(fields.values())
        values.append(reminder_id)
        
        conn = sqlite3.connect('reminders.db')
        cursor = conn.cursor()
        
        # Update the reminder
        cursor.execute(f'''
        UPDATE reminders
        SET {set_clause}
        WHERE id = ?
        ''', values)
        
        if cursor.rowcount == 0:
            conn.close()
            return jsonify({"error": "Reminder not found"}), 404
        
        conn.commit()
        conn.close()
        
        return jsonify({
            "success": True, 
            "message": "Reminder updated successfully"
        })
    
    except Exception as e:
        print(f"Error updating reminder: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/delete_reminder', methods=['POST'])
def delete_reminder():
    try:
        data = request.json
        reminder_id = data.get('id')
        
        if not reminder_id:
            return jsonify({"error": "Reminder ID is required"}), 400
        
        conn = sqlite3.connect('reminders.db')
        cursor = conn.cursor()
        
        # Instead of actually deleting, we just set active to 0
        cursor.execute('''
        UPDATE reminders
        SET active = 0
        WHERE id = ?
        ''', (reminder_id,))
        
        if cursor.rowcount == 0:
            conn.close()
            return jsonify({"error": "Reminder not found"}), 404
        
        conn.commit()
        conn.close()
        
        return jsonify({
            "success": True, 
            "message": "Reminder deleted successfully"
        })
    
    except Exception as e:
        print(f"Error deleting reminder: {str(e)}")
        return jsonify({"error": str(e)}), 500


# API endpoint for LLM queries
@app.route('/query', methods=['POST'])
def process_query():
    global qa_chain
    
    if not qa_chain:
        success = initialize_qa_system()
        if not success:
            return jsonify({"error": "Failed to initialize the QA system. Please check logs."})
    
    data = request.json
    user_query = data.get('query', '')
    
    if not user_query:
        return jsonify({"error": "Empty query"})
    
    try:
        # Format the query with system message
        formatted_query = f"{SYSTEM_MESSAGE}\n\nUser Query: {user_query}"
        
        # Process the query through the QA chain
        response = qa_chain.invoke(formatted_query)
        return jsonify({"response": response["result"]})
    except Exception as e:
        print(f"Error processing query: {e}")
        return jsonify({"error": f"Error processing query: {str(e)}"})

if __name__ == "__main__":
    # Initialize the QA system when the app starts
    initialize_qa_system()
    init_app()
    # Run the Flask app
    app.run(debug=True, host='0.0.0.0', port=5000)