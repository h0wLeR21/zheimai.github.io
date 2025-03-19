import os
import random
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class AlzheimersAssessment:
    def __init__(self, patient_name):
        self.patient_name = patient_name
        self.all_questions = []
        self.selected_questions = []
        self.answers = []
        self.assessment_date = datetime.now()
        
    def load_questions(self, filename="questions.docx", num_questions=10):
        """Load questions from a Word document and select a random subset."""
        if not os.path.exists(filename):
            raise FileNotFoundError(f"Questions file '{filename}' not found. Please create a questions.docx file with your assessment questions.")
        
        doc = Document(filename)
        self.all_questions = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        
        # Select random questions
        if len(self.all_questions) < num_questions:
            self.selected_questions = self.all_questions  # Use all if we have fewer than requested
            print(f"Warning: Only {len(self.all_questions)} questions available. All will be used.")
        else:
            self.selected_questions = random.sample(self.all_questions, num_questions)
        
    def conduct_assessment(self):
        """Conduct the assessment by asking questions and recording answers."""
        print(f"\nBeginning assessment for {self.patient_name}")
        print("=" * 50)
        
        for i, question in enumerate(self.selected_questions, 1):
            print(f"\nQuestion {i}: {question}")
            answer = input("Answer: ")
            self.answers.append(answer)
            
    def get_next_index(self, base_filename):
        """Get the next available index for the report filename."""
        index = 1
        while True:
            if not os.path.exists(f"{base_filename}_{index}.docx"):
                return index
            index += 1
            
    def generate_report(self, output_filename=None):
        """Generate a detailed assessment report."""
        reports_dir = "Reports"
        os.makedirs(reports_dir, exist_ok=True)  # Ensure "Reports" directory exists

        if output_filename is None:
            date_str = self.assessment_date.strftime("%Y%m%d")
            base_filename = f"assessment_report_{self.patient_name}_{date_str}"
            index = self.get_next_index(os.path.join(reports_dir, base_filename))
            output_filename = os.path.join(reports_dir, f"{base_filename}_{index}.docx")
        
        doc = Document()
        
        # Add header
        doc.add_heading('Memory Assessment Report', 0)
        
        # Add patient information
        doc.add_paragraph(f'Patient Name: {self.patient_name}')
        doc.add_paragraph(f'Assessment Date: {self.assessment_date.strftime("%B %d, %Y")}')
        doc.add_paragraph(f'Assessment Time: {self.assessment_date.strftime("%I:%M %p")}')
        
        # Get assessment number without f-string syntax error
        date_str = self.assessment_date.strftime("%Y%m%d")
        base_path = os.path.join(reports_dir, f"assessment_report_{self.patient_name}_{date_str}")
        assessment_number = self.get_next_index(base_path)
        doc.add_paragraph(f'Assessment Number: {assessment_number}')
        
        # Add questions and answers
        doc.add_heading('Assessment Details', level=1)
        for question, answer in zip(self.selected_questions, self.answers):
            p = doc.add_paragraph()
            p.add_run('Question: ').bold = True
            p.add_run(question)
            p.add_run('\nAnswer: ').bold = True
            p.add_run(answer)
            doc.add_paragraph()  # Add spacing between Q&A pairs
            
        # Add observations section
        doc.add_heading('Clinical Observations', level=1)
        doc.add_paragraph('Memory Performance Observations:')
        doc.add_paragraph('• Response Time and Clarity:')
        doc.add_paragraph('• Behavioral Observations:')
        doc.add_paragraph('• Additional Notes:')
        
        # Add recommendations section
        doc.add_heading('Recommendations', level=1)
        doc.add_paragraph('Based on today\'s assessment:')
        
        # Save the report
        doc.save(output_filename)
        return output_filename

def main():
    # Example usage
    patient_name = input("Enter patient name: ")
    
    # Create assessment instance
    assessment = AlzheimersAssessment(patient_name)
    
    try:
        # Load questions from file and select 10 random ones
        assessment.load_questions(num_questions=10)
        
        # Conduct the assessment
        assessment.conduct_assessment()
        
        # Generate and save report
        report_file = assessment.generate_report()
        print(f"\nAssessment complete. Report saved as: {report_file}")
        
    except FileNotFoundError as e:
        print(f"Error: {e}")
    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    main()